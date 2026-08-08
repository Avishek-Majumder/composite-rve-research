#!/usr/bin/env python3

"""
Generate the formal post-M5 Secondary Planning document.

This document preserves the original Project 01 plan while recording the
agreed roadmap correction made after completion of M5.

The generated DOCX is a project governance and research-planning record.
It does not execute FEM simulations and does not begin M6.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUTPUT_PATH = Path("docs/Secondary_Planning.docx")

DOCUMENT_TITLE = "Secondary Planning"
DOCUMENT_SUBTITLE = (
    "Post-M5 Corrected Research Roadmap and Project Governance Record"
)

PROJECT_DIRECTION = (
    "An Active-Learning and Finite-Element Framework for "
    "Uncertainty-Aware Prediction of Defect-Sensitive Mechanical "
    "Properties in Particle-Reinforced Composites."
)

REVISION_LABEL = "Post-M5 alignment revision"
REVISION_DATE = "9 August 2026"

M5_CLOSURE_COMMIT = (
    "5914bf50108c042019303ece09552768faeeb977"
)

PRE_DOCUMENTATION_COMMIT = (
    "05c5fc7e9d7315c6aaa48af78c9290319f1a0b6b"
)

D0_SHA256 = (
    "e48e9eb731b6e13eb15b33ab643722a9d72e8bdb933bf24d9c1c3847776c17d1"
)


MILESTONES_COMPLETED = [
    (
        "M0",
        "Software and development environment",
        "100% COMPLETE",
    ),
    (
        "M1",
        "FEM fundamentals and homogeneous validation",
        "100% COMPLETE",
    ),
    (
        "M2",
        "First single-particle composite model",
        "100% COMPLETE",
    ),
    (
        "M3",
        "Mesh Convergence and Composite Verification",
        "100% COMPLETE",
    ),
    (
        "M4",
        "Parametric RVE and sampling foundation",
        "100% COMPLETE",
    ),
    (
        "M5",
        "Initial perfect-bonding FEM dataset generation",
        "100% COMPLETE",
    ),
]


CORRECTED_ROADMAP = [
    (
        "M6",
        "Multiple/Random-Particle Microstructure Foundation",
        (
            "Develop a reproducible multi-particle geometry generator with "
            "stored random seeds, overlap prevention, minimum-spacing rules, "
            "particle-size variability, and random/clustered arrangements."
        ),
        (
            "Reproducible geometry generation and validity checks pass; "
            "invalid or unmeshable geometries are explicitly recorded rather "
            "than disappearing silently."
        ),
    ),
    (
        "M7",
        "Circular Void Defects and Defect-Sensitive Response Definition",
        (
            "Introduce circular voids as the Version-1 defect model and "
            "define a robust defect-sensitive stress-concentration output."
        ),
        (
            "Void geometries mesh and solve reliably; the defect-sensitive "
            "stress target is physically interpretable and its mesh "
            "sensitivity has been assessed."
        ),
    ),
    (
        "M8",
        (
            "RVE-Size Study, Homogenization BC/PBC Verification, "
            "and Final Target-Mesh Verification"
        ),
        (
            "Separate statistical RVE representativity from element-level "
            "mesh convergence, verify the final homogenization boundary "
            "conditions, and establish the production mesh policy for the "
            "stochastic defect simulations."
        ),
        (
            "RVE-size behavior is quantified across realizations; the final "
            "boundary-condition strategy is verified; the selected mesh is "
            "adequate for the final primary targets."
        ),
    ),
    (
        "M9",
        "Final Parameter-Space Lock and Stochastic Pilot Dataset",
        (
            "Lock the final microstructure/defect parameter space, seed "
            "policy, result schema, provenance fields, and execute a limited "
            "pilot before expensive production runs."
        ),
        (
            "Pilot runs demonstrate an end-to-end reproducible pipeline, "
            "failure reporting, runtime feasibility, and physically sensible "
            "parameter ranges."
        ),
    ),
    (
        "M10",
        "Main Quality-Controlled FEM Simulation Database",
        (
            "Generate the principal stochastic FEM database using multiple "
            "microstructure realizations per relevant parameter condition."
        ),
        (
            "The database is quality controlled, reproducible, provenance "
            "complete, and suitable for leakage-safe ML experiments."
        ),
    ),
    (
        "M11",
        "Baseline Machine-Learning Models and Grouped Validation",
        (
            "Train interpretable conventional surrogate models using grouped "
            "splits that prevent closely related microstructures from leaking "
            "between training and validation data."
        ),
        (
            "Baseline model comparison is reproducible and reported with "
            "appropriate regression metrics and physical sanity checks."
        ),
    ),
    (
        "M12",
        "Active Learning versus Random Sampling",
        (
            "Run uncertainty-guided active learning and compare it against "
            "random acquisition under equal FEM simulation budgets."
        ),
        (
            "The study can quantify whether active learning reduces "
            "prediction error or improves data efficiency relative to the "
            "random baseline."
        ),
    ),
    (
        "M13",
        "Uncertainty Calibration, Variability, and OOD Testing",
        (
            "Separate microstructure variability from model uncertainty, "
            "evaluate uncertainty calibration, and perform deliberate "
            "out-of-distribution tests."
        ),
        (
            "Prediction uncertainty is quantitatively evaluated rather than "
            "merely displayed, and unfamiliar regimes are explicitly tested."
        ),
    ),
    (
        "M14",
        "Final Analysis, Ablations, Figures, and Manuscript",
        (
            "Consolidate scientific findings, limitations, ablations, final "
            "figures/tables, reproducibility records, and manuscript-ready "
            "results."
        ),
        (
            "The final claims are supported by traceable simulations, "
            "validated ML experiments, uncertainty evidence, and a clean "
            "reproducible repository."
        ),
    ),
]


OUTSTANDING_REQUIREMENTS = [
    (
        "Multiple/random-particle microstructures",
        (
            "The present FEM baseline uses one centered circular particle. "
            "The final research problem still requires stochastic "
            "multi-particle configurations."
        ),
    ),
    (
        "Circular void defects",
        (
            "No defect is present in D0-PB. Circular voids remain the planned "
            "Version-1 defect representation."
        ),
    ),
    (
        "RVE-size/statistical representativity",
        (
            "M3 established mesh convergence, but mesh convergence and RVE "
            "representativity answer different questions. A stochastic RVE "
            "size study is still required."
        ),
    ),
    (
        "Final homogenization boundary conditions",
        (
            "The existing displacement boundary conditions remain valid for "
            "the verified baseline problem, but random-position production "
            "data must not be generated at scale until the final "
            "homogenization BC/PBC strategy has been verified."
        ),
    ),
    (
        "Defect-sensitive stress target",
        (
            "Raw local stress extrema remain diagnostics only. A more robust "
            "stress-concentration indicator must be defined and mesh checked "
            "before becoming a primary ML response."
        ),
    ),
    (
        "Repeated microstructure realizations",
        (
            "Multiple seeds are required to distinguish variation caused by "
            "microstructure realization from uncertainty caused by the "
            "surrogate model."
        ),
    ),
    (
        "Final provenance and failure schema",
        (
            "The production stochastic dataset should preserve case ID, "
            "physical parameters, random seed, geometry status, mesh status, "
            "solver status, element count, runtime, failure reason, and "
            "output provenance."
        ),
    ),
    (
        "Broader physics validation",
        (
            "Random-RVE convergence, void-related physical trend checks, "
            "appropriate theoretical comparisons, and selected published "
            "numerical comparisons remain future validation tasks."
        ),
    ),
]


LOCKED_BASELINE_ASSUMPTIONS = [
    "Two-dimensional computational representation.",
    "Small-strain linear elasticity.",
    "Plane-stress formulation through M5.",
    "Isotropic matrix and isotropic reinforcing particle materials.",
    "Perfect matrix-particle bonding through M5.",
    "Normalized RVE width = 1.0 and height = 1.0.",
    "Centered particle for the D0-PB baseline.",
    "Matrix Young's modulus = 1000.0 for the D0-PB baseline.",
    "Prescribed x displacement = 0.01 for the D0-PB baseline.",
    "Production mesh used for D0-PB = 0.02048.",
    (
        "Future changes such as stochastic particle placement, circular "
        "voids, or PBCs must occur only inside their explicitly announced "
        "milestones and must be separately verified."
    ),
]


D0_FACTS = [
    ("Formal interpretation", "D0-PB baseline dataset"),
    (
        "Geometry class",
        "Centered single circular inclusion; no defect",
    ),
    ("Interface", "Perfect bonding"),
    ("Cases", "60/60 successful"),
    ("Failed FEM cases", "0"),
    ("Aggregate rows", "60"),
    ("Aggregate columns", "45"),
    (
        "Effective modulus range",
        "1051.948603727062 to 1589.3874813628656",
    ),
    (
        "Effective Poisson response range",
        "0.21573865543142967 to 0.3863453188926385",
    ),
    (
        "Aggregate file",
        "results/processed/07_m5_initial_fem_dataset.csv",
    ),
    ("Aggregate SHA-256", D0_SHA256),
]


ML_GUARDRAILS = [
    (
        "Do not treat D0-PB as the final defect-sensitive research dataset."
    ),
    (
        "Do not begin final surrogate claims before the stochastic "
        "microstructure and defect physics are implemented and validated."
    ),
    (
        "Use grouped/leakage-safe validation once multiple related "
        "microstructure realizations exist."
    ),
    (
        "Begin with conventional tabular baselines rather than neural "
        "networks unless later evidence justifies additional complexity."
    ),
    (
        "Compare active learning with random selection using equal FEM "
        "simulation budgets."
    ),
    (
        "Evaluate uncertainty calibration quantitatively and include "
        "deliberate out-of-distribution tests."
    ),
]


GOVERNANCE_RULES = [
    (
        "Original Project 01 remains the initial project plan and historical "
        "scope record."
    ),
    (
        "Secondary Planning is the authoritative post-M5 roadmap correction; "
        "it supplements rather than deletes the original plan."
    ),
    (
        "M0-M5 remain completed and are not renumbered or repeated."
    ),
    (
        "Every major milestone transition requires an explicit completion "
        "summary, next-milestone goal, progress state, and user confirmation."
    ),
    (
        "Research work continues strictly one concrete terminal step at a "
        "time, with output verification before continuation."
    ),
    (
        "For every meaningful validated repository unit: inspect changes, "
        "stage only intended paths, run staged safety checks, commit, push "
        "to origin/main, verify HEAD equals origin/main, and finish clean."
    ),
    (
        "PROJECT_STATUS.md should be maintained as a concise repository "
        "checkpoint so roadmap state is recoverable independently of chat "
        "memory."
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, rgb: RGBColor) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb


def set_cell_bold(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    """Keep a table row together on one page when possible."""

    tr_pr = row._tr.get_or_add_trPr()

    existing = tr_pr.find(
        qn("w:cantSplit")
    )

    if existing is None:
        cant_split = OxmlElement(
            "w:cantSplit"
        )
        tr_pr.append(cant_split)


def format_table(table, header_fill: str = "1F4E78") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    if not table.rows:
        return

    header = table.rows[0]
    set_repeat_table_header(header)

    for cell in header.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_bold(cell)
        set_cell_text_color(
            cell,
            RGBColor(255, 255, 255),
        )

    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_table_pagination_controls(
    document: Document,
) -> None:
    """Prevent completed table rows from splitting across pages."""

    for table in document.tables:
        for row in table.rows:
            prevent_row_split(row)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(
        text,
        style="List Bullet",
    )
    paragraph.paragraph_format.space_after = Pt(3)


def add_numbered(
    document: Document,
    number: int,
    text: str,
) -> None:
    """Add an explicit independently numbered planning item."""

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = Mm(7)
    paragraph.paragraph_format.first_line_indent = Mm(-7)
    paragraph.paragraph_format.space_after = Pt(3)

    paragraph.add_run(
        f"{number}.  {text}"
    )


def configure_document(document: Document) -> None:
    section = document.sections[0]

    section.page_width = Mm(210)
    section.page_height = Mm(297)

    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size in [
        ("Title", 24),
        ("Heading 1", 16),
        ("Heading 2", 12.5),
        ("Heading 3", 11),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 120)

    document.styles["Heading 1"].paragraph_format.space_before = Pt(10)
    document.styles["Heading 1"].paragraph_format.space_after = Pt(6)

    document.styles["Heading 2"].paragraph_format.space_before = Pt(8)
    document.styles["Heading 2"].paragraph_format.space_after = Pt(4)

    document.styles["Normal"].paragraph_format.space_after = Pt(5)
    document.styles["Normal"].paragraph_format.line_spacing = 1.08

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = (
        "Secondary Planning | Post-M5 Corrected Roadmap | "
        "Simulation + ML, no laboratory experiments"
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(80)

    run = paragraph.add_run(DOCUMENT_TITLE)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 120)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subtitle.add_run(DOCUMENT_SUBTITLE)
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(70, 70, 70)

    document.add_paragraph()

    project = document.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = project.add_run(PROJECT_DIRECTION)
    run.bold = True
    run.font.size = Pt(11.5)

    document.add_paragraph()

    details = document.add_table(rows=4, cols=2)
    details.style = "Table Grid"
    details.alignment = WD_TABLE_ALIGNMENT.CENTER

    items = [
        ("Research route", "Simulation + Machine Learning only"),
        ("Laboratory experiments", "None"),
        ("Revision", REVISION_LABEL),
        ("Revision date", REVISION_DATE),
    ]

    for row, (label, value) in zip(details.rows, items):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bold(row.cells[0])
        set_cell_shading(row.cells[0], "D9EAF7")

    document.add_paragraph()

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = note.add_run(
        "Status: Approved planning correction after the formal closure of M5."
    )
    run.bold = True

    document.add_page_break()


def add_executive_decision(document: Document) -> None:
    document.add_heading("1. Executive Decision", level=1)

    document.add_paragraph(
        (
            "The project remains scientifically viable and the work completed "
            "through M5 is retained as valid. The post-M5 alignment audit "
            "identified a roadmap-order drift rather than invalid research. "
            "M4 and M5 are therefore preserved as useful inserted foundation "
            "milestones."
        )
    )

    document.add_paragraph(
        (
            "The project will NOT proceed directly from M5 into machine "
            "learning. The missing random-microstructure, void-defect, RVE, "
            "homogenization-boundary-condition, stochastic-variability, and "
            "final dataset stages must be completed first."
        )
    )

    decision = document.add_table(rows=1, cols=2)
    decision.rows[0].cells[0].text = "Decision"
    decision.rows[0].cells[1].text = "Approved interpretation"
    format_table(decision)

    rows = [
        (
            "M0-M5",
            "Remain completed; no renumbering or repetition.",
        ),
        (
            "M4-M5",
            (
                "Inserted parametric/batch-pipeline foundation; they do not "
                "replace random-microstructure or defect milestones."
            ),
        ),
        (
            "M5 dataset",
            (
                "Reclassified conceptually as D0-PB, a deterministic "
                "perfect-bonding centered-single-inclusion baseline."
            ),
        ),
        (
            "Next scientific milestone",
            "M6 - Multiple/Random-Particle Microstructure Foundation.",
        ),
        (
            "ML readiness",
            (
                "Final research ML is not ready to start until the corrected "
                "physics/data milestones are completed."
            ),
        ),
    ]

    for item, value in rows:
        cells = decision.add_row().cells
        cells[0].text = item
        cells[1].text = value

    document.add_paragraph(
        (
            "This Secondary Planning record supplements the original Project "
            "01 plan. The original plan remains the historical scope record."
        )
    )


def add_completed_state(document: Document) -> None:
    document.add_heading("2. Validated State Through M5", level=1)

    table = document.add_table(rows=1, cols=3)

    for index, text in enumerate(
        ["Milestone", "Description", "Status"]
    ):
        table.rows[0].cells[index].text = text

    format_table(table)

    for milestone, description, status in MILESTONES_COMPLETED:
        cells = table.add_row().cells
        cells[0].text = milestone
        cells[1].text = description
        cells[2].text = status

    document.add_heading("2.1 M5 closure record", level=2)

    for text in [
        (
            "M5 produced a validated batch-execution and deterministic "
            "aggregation workflow for the locked 60-case perfect-bonding "
            "Latin-hypercube design."
        ),
        "All 60 FEM cases completed successfully; zero failed cases.",
        "All FEM verification suites passed for the 60 cases.",
        (
            "Raw per-case JSON records and runtime logs remain local runtime "
            "artifacts and are ignored by Git."
        ),
        (
            "The validated aggregate is tracked in Git and was reproduced "
            "byte-for-byte before M5 closure."
        ),
        f"M5 closure commit: {M5_CLOSURE_COMMIT}",
        (
            "Documentation-dependency checkpoint before generation of this "
            f"record: {PRE_DOCUMENTATION_COMMIT}"
        ),
    ]:
        add_bullet(document, text)


def add_d0_interpretation(document: Document) -> None:
    document.add_heading(
        "3. D0-PB Baseline Dataset Interpretation",
        level=1,
    )

    document.add_paragraph(
        (
            "The file results/processed/07_m5_initial_fem_dataset.csv is "
            "retained as a valuable research artifact, but it must not be "
            "treated as the final defect-sensitive stochastic dataset."
        )
    )

    document.add_paragraph(
        (
            "Its formal role from this planning revision onward is D0-PB: "
            "the validated perfect-bonding, centered-single-inclusion "
            "parametric baseline dataset."
        )
    )

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Property"
    table.rows[0].cells[1].text = "Validated value"
    format_table(table)

    for label, value in D0_FACTS:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_heading("3.1 Permitted uses of D0-PB", level=2)

    for item in [
        "Regression testing of the FEM-to-dataset pipeline.",
        "Baseline surrogate-method development and debugging.",
        "Comparison against later stochastic/defect microstructures.",
        "Verification of data schemas, preprocessing, and ML infrastructure.",
        (
            "A deterministic reference demonstrating that parameterized "
            "perfect-bonding FEM automation works."
        ),
    ]:
        add_bullet(document, item)

    document.add_heading("3.2 Prohibited interpretation", level=2)

    document.add_paragraph(
        (
            "D0-PB must not be used to support final claims about void "
            "defects, random particle arrangements, microstructure "
            "variability, OOD behavior, uncertainty reliability, or active "
            "learning efficiency."
        )
    )


def add_outstanding_requirements(document: Document) -> None:
    document.add_heading(
        "4. Scientific Work Still Required Before Final ML",
        level=1,
    )

    for index, (title, description) in enumerate(
        OUTSTANDING_REQUIREMENTS,
        start=1,
    ):
        document.add_heading(
            f"4.{index} {title}",
            level=2,
        )
        document.add_paragraph(description)


def add_corrected_roadmap(document: Document) -> None:
    document.add_heading(
        "5. Corrected Future Roadmap: M6-M14",
        level=1,
    )

    document.add_paragraph(
        (
            "The following roadmap is the accepted continuation path. "
            "Existing milestone numbers M0-M5 remain unchanged."
        )
    )

    table = document.add_table(rows=1, cols=4)

    headers = [
        "Milestone",
        "Name",
        "Primary purpose",
        "Completion gate",
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    format_table(table)

    for milestone, name, purpose, gate in CORRECTED_ROADMAP:
        cells = table.add_row().cells
        cells[0].text = milestone
        cells[1].text = name
        cells[2].text = purpose
        cells[3].text = gate


def add_assumptions(document: Document) -> None:
    document.add_heading(
        "6. Baseline Physics and Terminology Guardrails",
        level=1,
    )

    document.add_paragraph(
        (
            "The assumptions below describe the validated baseline through "
            "M5. They must not be silently altered."
        )
    )

    for item in LOCKED_BASELINE_ASSUMPTIONS:
        add_bullet(document, item)

    document.add_heading("6.1 2D fraction terminology", level=2)

    document.add_paragraph(
        (
            "For this two-dimensional model, use 'particle area fraction' "
            "or '2D particle fraction' when describing the simulated "
            "reinforcement fraction. Avoid presenting it as a directly "
            "measured three-dimensional particle volume fraction."
        )
    )

    document.add_heading("6.2 Local stresses", level=2)

    document.add_paragraph(
        (
            "Local stress extrema remain diagnostic quantities. M3 showed "
            "that some extrema are more mesh sensitive than the global "
            "effective properties. They are not automatically accepted as "
            "final ML targets."
        )
    )


def add_ml_guardrails(document: Document) -> None:
    document.add_heading(
        "7. ML, Active-Learning, and Uncertainty Guardrails",
        level=1,
    )

    for number, item in enumerate(
        ML_GUARDRAILS,
        start=1,
    ):
        add_numbered(
            document,
            number,
            item,
        )

    document.add_heading(
        "7.1 Planned primary effective-property targets",
        level=2,
    )

    for item in [
        "Effective axial modulus.",
        "Effective Poisson response.",
        (
            "A robust defect-sensitive stress-concentration indicator, "
            "only after M7 defines and verifies it."
        ),
    ]:
        add_bullet(document, item)


def add_governance(document: Document) -> None:
    document.add_heading(
        "8. Repository and Milestone Governance",
        level=1,
    )

    for number, item in enumerate(
        GOVERNANCE_RULES,
        start=1,
    ):
        add_numbered(
            document,
            number,
            item,
        )

    document.add_heading(
        "8.1 Required persistent project records",
        level=2,
    )

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Record"
    table.rows[0].cells[1].text = "Purpose"
    format_table(table)

    records = [
        (
            "Original Project 01 planning",
            "Historical initial scope and research intent.",
        ),
        (
            "docs/Secondary_Planning.docx",
            (
                "Formal approved post-M5 roadmap correction and scientific "
                "planning record."
            ),
        ),
        (
            "PROJECT_STATUS.md",
            (
                "Concise current milestone, latest validated checkpoint, "
                "next step, assumptions, and Git synchronization state."
            ),
        ),
        (
            "Git/GitHub history",
            (
                "Version-controlled evidence of validated implementation and "
                "planning decisions."
            ),
        ),
    ]

    for record, purpose in records:
        cells = table.add_row().cells
        cells[0].text = record
        cells[1].text = purpose


def add_next_action(document: Document) -> None:
    document.add_heading(
        "9. Immediate Project State and Next Gate",
        level=1,
    )

    document.add_paragraph(
        (
            "At the time of this Secondary Planning revision, M0-M5 are "
            "formally complete and M6 has not started."
        )
    )

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "State"
    format_table(table)

    states = [
        ("M0-M5", "100% COMPLETE"),
        ("Project-plan alignment audit", "COMPLETE"),
        (
            "Secondary Planning roadmap",
            "APPROVED; repository documentation in progress",
        ),
        (
            "Final defect-sensitive stochastic dataset",
            "NOT YET GENERATED",
        ),
        ("Random microstructures", "NOT YET IMPLEMENTED"),
        ("Circular void defects", "NOT YET IMPLEMENTED"),
        ("Final research ML phase", "NOT READY TO START"),
        (
            "Next major milestone",
            "M6 - Multiple/Random-Particle Microstructure Foundation",
        ),
    ]

    for item, state in states:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = state

    document.add_paragraph()

    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        (
            "M6 must begin only after this planning/documentation checkpoint "
            "has been generated, validated, committed, pushed, and the user "
            "explicitly confirms the M6 transition."
        )
    )
    run.bold = True


def add_source_basis(document: Document) -> None:
    document.add_heading(
        "10. Source Basis for This Secondary Plan",
        level=1,
    )

    document.add_paragraph(
        (
            "This planning record was prepared by reconciling the original "
            "Project 01 research plan with the validated implementation "
            "history through M5, the Project 02-04 continuation records, the "
            "repository state, and the explicitly approved post-M5 alignment "
            "discussion."
        )
    )

    document.add_paragraph(
        (
            "Where the execution order differed from the original plan, the "
            "completed work was preserved when scientifically valid and the "
            "missing original-scope requirements were restored to the future "
            "roadmap rather than silently discarded."
        )
    )


def build_document() -> Document:
    document = Document()
    configure_document(document)

    add_title_page(document)
    add_executive_decision(document)
    add_completed_state(document)
    add_d0_interpretation(document)
    add_outstanding_requirements(document)
    add_corrected_roadmap(document)
    add_assumptions(document)
    add_ml_guardrails(document)
    add_governance(document)
    add_next_action(document)
    add_source_basis(document)

    # Apply pagination controls only after every table and row exists.
    apply_table_pagination_controls(document)

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = build_document()
    document.save(OUTPUT_PATH)

    print(
        "Secondary Planning DOCX generated:",
        OUTPUT_PATH,
    )


if __name__ == "__main__":
    main()
